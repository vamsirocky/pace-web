# server/org_app/router.py
from fastapi import APIRouter, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import text
from typing import Optional, Dict, Tuple, List
from datetime import datetime
import io
import os
from openai import OpenAIError
# LLM + DOCX
from openai import OpenAI
from docx import Document
import os, requests
from pathlib import Path
from dotenv import load_dotenv

from ..qlearning_app.db_store_pg import engine  

router = APIRouter()

CAT_PREFIX: Dict[str, str] = {
    "donate": "CAT01_%",
    "volunteer": "CAT02_%",
    "advocate": "CAT03_%",
    "wellness": "CAT04_%",
    "recycle": "CAT05_%",
    "wildlife": "CAT06_%",
}


ACTION_LABELS: Dict[str, str] = {
    "CAT01_A1": "Note Sharing Day",
    "CAT01_A2": "Buy eco product",
    "CAT02_A1": "One-Tap Survey",
    "CAT02_A2": "Lead a drive",
    "CAT03_A1": "Selfie Spot",
    "CAT03_A2": "Startup Pitch",
    "CAT04_A1": "Stairs Challenge",
    "CAT04_A2": "Gym Power-Up",
    "CAT05_A1": "Refill & Reuse",
    "CAT05_A2": "Reusable Mug",
    "CAT06_A1": "Meatless Monday",
    "CAT06_A2": "Report a sighting",
}

def _users_meta(conn) -> Dict[str, str]:
    """Return {'column_name': 'data_type', ...} for public.users (if present)."""
    rows = conn.execute(
        text("""
            SELECT column_name, data_type
            FROM information_schema.columns
            WHERE table_schema='public' AND table_name='users'
        """)
    ).fetchall()
    return {r[0]: r[1] for r in rows}

def _build_user_join_and_display(conn) -> Tuple[str, str]:
    """
    Build a safe LEFT JOIN to public.users and a display expr for the name.
    We always compare as TEXT to avoid uuid/text operator errors.
    """
    cols = _users_meta(conn)
    if not cols:
        return "", "al.user_id"

    # picks a join key in users table
    join_key = next((c for c in ("user_id", "id", "uid", "uuid", "auth_id", "email") if c in cols), None)

    # picks best name/display expression
    if "first_name" in cols and "last_name" in cols:
        name_expr = "NULLIF(trim(coalesce(u.first_name,'') || ' ' || coalesce(u.last_name,'')), '')"
    elif "name" in cols:
        name_expr = "u.name"
    elif "full_name" in cols:
        name_expr = "u.full_name"
    elif "display_name" in cols:
        name_expr = "u.display_name"
    elif "username" in cols:
        name_expr = "u.username"
    elif "email" in cols:
        name_expr = "u.email"
    else:
        name_expr = None

    if join_key:
        # Cast users key to text to match activity_log.user_id (TEXT)
        join_clause = f"LEFT JOIN public.users u ON u.{join_key}::text = al.user_id"
        display = f"COALESCE(({name_expr})::text, al.user_id)" if name_expr else "al.user_id"
    else:
        join_clause = ""
        display = "al.user_id"

    return join_clause, display

def _where_time_and_category(days: int, category: Optional[str]) -> Tuple[str, Dict]:
    """Return SQL WHERE filters for time window and category prefix, with params."""
    where = ["al.ts >= now() - make_interval(days => :days)"]
    params: Dict = {"days": int(days)}
    if category and category in CAT_PREFIX:
        where.append("al.action LIKE :prefix")
        params["prefix"] = CAT_PREFIX[category]
    return " AND ".join(where), params

# ----------------- Existing endpoints  -----------------

@router.get("/metrics")
def metrics(days: int = 30, bucket: str = Query("day", pattern="^(day|week|month)$"), category: Optional[str] = None):
    """
    KPI totals for the time window (optionally filtered by category).
    participants: distinct users
    completions: count of rows
    pointsTotal: sum(reward)
    rate: completions/participants * 100 (rounded)
    """
    with engine.begin() as conn:
        where_sql, params = _where_time_and_category(days, category)
        sql = f"""
            WITH filtered AS (
              SELECT al.user_id, al.reward
              FROM public.activity_log al
              WHERE {where_sql}
            )
            SELECT
              COUNT(DISTINCT user_id) AS participants,
              COUNT(*) AS completions,
              COALESCE(SUM(reward), 0) AS points_total
            FROM filtered
        """
        row = conn.execute(text(sql), params).one()
        participants = int(row[0] or 0)
        completions = int(row[1] or 0)
        points_total = float(row[2] or 0.0)
        rate = round((completions / participants) * 100) if participants else 0

    return {
        "participants": participants,
        "completions": completions,
        "pointsTotal": points_total,
        "rate": rate,
    }

@router.get("/actions")
def actions(days: int = 30, bucket: str = Query("day", pattern="^(day|week|month)$"),
            category: Optional[str] = None, limit: int = 12):
    """
    Top actions by completions in the window. Always returns a 2-point trend.
    """
    with engine.begin() as conn:
        where_sql, params = _where_time_and_category(days, category)
        params["limit"] = int(limit)
        sql = f"""
            SELECT al.action AS id, COUNT(*)::int AS completions
            FROM public.activity_log al
            WHERE {where_sql}
            GROUP BY al.action
            ORDER BY completions DESC
            LIMIT :limit
        """
        rows = conn.execute(text(sql), params).mappings().all()

   
    out = []
    for r in rows:
        comp = int(r["completions"] or 0)
        trend = [max(0, comp // 2), comp]  # at least 2 points
        out.append({
            "id": r["id"],
            "label": r["id"],       
            "completions": comp,
            "trend": trend,
        })
    return out

@router.get("/recent")
def recent(limit: int = 20, category: Optional[str] = None):
    """
    Most recent verified activity (optionally filtered by category).
    Joins users safely (users key ::text = al.user_id) and shows a friendly name.
    """
    with engine.begin() as conn:
        join_clause, display_expr = _build_user_join_and_display(conn)
        where = ["1=1"]
        params: Dict = {"limit": int(limit)}

        if category and category in CAT_PREFIX:
            where.append("al.action LIKE :prefix")
            params["prefix"] = CAT_PREFIX[category]

        sql = f"""
            SELECT
              to_char(al.ts, 'YYYY-MM-DD HH24:MI') AS ts,
              {display_expr} AS user_display,
              al.action,
              al.reward
            FROM public.activity_log al
            {join_clause}
            WHERE {' AND '.join(where)}
            ORDER BY al.ts DESC
            LIMIT :limit
        """
        rows = conn.execute(text(sql), params).mappings().all()

    # derive category on client
    return [
        {
            "ts": r["ts"],
            "user": r["user_display"],
            "action": r["action"],
            "points": float(r["reward"]),
        }
        for r in rows
    ]

# ----------------- NEW: SDG/ESG report as a .docx download -----------------

def _fetch_activity_rollup(days: int, category: Optional[str]) -> List[Dict]:
    """List of {action,label,persons,contributions} for the window."""
    with engine.begin() as conn:
        where_sql, params = _where_time_and_category(days, category)
        sql = f"""
            SELECT
              al.action,
              COUNT(DISTINCT al.user_id)::int AS persons,
              COALESCE(SUM(al.reward), 0)::float AS contributions
            FROM public.activity_log al
            WHERE {where_sql}
            GROUP BY al.action
            ORDER BY persons DESC, contributions DESC
        """
        rows = conn.execute(text(sql), params).mappings().all()

    out = []
    for r in rows:
        aid = r["action"]
        out.append({
            "action": aid,
            "label": ACTION_LABELS.get(aid, aid),
            "persons": int(r["persons"] or 0),
            "contributions": float(r["contributions"] or 0.0),
        })
    return out

def _build_prompt(company: str, activities: List[Dict]) -> str:
    # compact table the model can reason over
    header = "Activity | Description | Persons | Contributions\n--|--|--:|--:\n"
    lines = [f"{a['label']} | {a['action']} | {a['persons']} | {round(a['contributions'])}" for a in activities]
    table = header + ("\n".join(lines) if lines else "_No activity data_|_|0|0")

    return f"""
You are an ESG/SDG reporting analyst.

Organization: {company}

Recent activities:
{table}

Write a concise ESG/SDG report with these sections:
1) Executive Summary
2) ESG Strategy & Vision (inferred from activities)
3) Environmental contributions (analysis + gaps + suggestions)
4) Social contributions (analysis + gaps + suggestions)
5) Governance contributions (analysis + gaps + suggestions)
6) ESG Performance Metrics Summary (state assumptions; benchmark vs typical org)
7) SDG Mapping (which UN SDGs the activities align with, and why)
8) Next Steps (prioritized actions with short rationale)

Use clear headings and concise professional prose.
"""

OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")

def _call_llm(prompt: str, model: str | None = None) -> str | None:
    """
    Calls the selected LLM provider based on env LLM_PROVIDER.
    Supported: 'ollama' (local, free), 'huggingface', 'openrouter', 'openai'.
    Returns a string or None on error.
    """
    provider = (os.getenv("LLM_PROVIDER") or "openai").lower()

    try:
        # ---- 1) OLLAMA (local, free) ----
        if provider == "ollama":
            host = os.getenv("OLLAMA_HOST", "http://127.0.0.1:11434")
            mdl = model or os.getenv("OLLAMA_MODEL", "llama3.1:8b")
            r = requests.post(
                f"{host}/api/generate",
                json={"model": mdl, "prompt": prompt, "stream": False},
                timeout=120,
            )
            r.raise_for_status()
            data = r.json()
            text = (data.get("response") or "").strip()
            return text or None

        # ---- 2) HUGGING FACE INFERENCE API (limited free tier) ----
        if provider == "huggingface":
            key = os.getenv("HF_API_TOKEN")
            if not key:
                return None
            mdl = model or os.getenv("HF_MODEL", "HuggingFaceH4/zephyr-7b-beta")
            r = requests.post(
                f"https://api-inference.huggingface.co/models/{mdl}",
                headers={"Authorization": f"Bearer {key}"},
                json={"inputs": prompt, "parameters": {"max_new_tokens": 600, "temperature": 0.3}},
                timeout=120,
            )
            r.raise_for_status()
            data = r.json()
            if isinstance(data, list):
                text = (data[0].get("generated_text") or data[0].get("summary_text") or "").strip()
            else:
                text = (data.get("generated_text") or data.get("summary_text") or "").strip()
            return text or None

        # ---- 3) OPENROUTER (aggregator; sometimes has free routes) ----
        if provider == "openrouter":
            key = os.getenv("OPENROUTER_API_KEY")
            if not key:
                return None
            mdl = model or os.getenv("OPENROUTER_MODEL", "meta-llama/llama-3.1-8b-instruct")
            headers = {
                "Authorization": f"Bearer {key}",
                # these two headers are recommended by OpenRouter
                "HTTP-Referer": os.getenv("APP_URL", "http://localhost:5173"),
                "X-Title": os.getenv("APP_NAME", "Pace Admin"),
                "Content-Type": "application/json",
            }
            payload = {
                "model": mdl,
                "messages": [
                    {"role": "system", "content": "You are an ESG/SDG reporting assistant."},
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.3,
            }
            r = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=payload, timeout=120)
            r.raise_for_status()
            data = r.json()
            text = (data.get("choices", [{}])[0].get("message", {}).get("content") or "").strip()
            return text or None

        # ---- 4) OPENAI (paid API) ----
        from openai import OpenAI
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        mdl = model or os.getenv("LLM_MODEL", "gpt-4o-mini")
        resp = client.chat.completions.create(
            model=mdl,
            messages=[
                {"role": "system", "content": "You are an ESG/SDG reporting assistant."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
        )
        text = (resp.choices[0].message.content or "").strip()
        return text or None

    except Exception as e:
        print(f"LLM error ({provider}):", e)
        return None

def _offline_narrative(org_name: str, period_label: str, metrics: dict, actions: list[dict]) -> str:
    """
    Build a simple narrative using local stats when LLM isn't available.
    """
    participants = int(metrics.get("participants", 0))
    completions = int(metrics.get("completions", 0))
    points = int(metrics.get("points_total", metrics.get("pointsTotal", 0)))
    rate = int(metrics.get("rate", 0))

    top = sorted(actions, key=lambda a: int(a.get("completions", 0)), reverse=True)[:3]
    top_str = ", ".join(f"{a.get('label', a.get('id'))} ({int(a.get('completions', 0))})" for a in top) if top else "n/a"

    lines = [
        f"This ESG/SDG report summarizes activity for {org_name} over {period_label}.",
        f"• {participants} unique participant(s) completed {completions} verified action(s) and earned {points} points.",
        f"• Completion rate (actions per participant) was approximately {rate}%.",
        f"• Top activities by completions: {top_str}.",
        "",
        "Impact Notes:",
        "• Participation breadth and frequency indicate engagement momentum.",
        "• Consider promoting under-represented categories to balance outcomes.",
        "• Next cycle, set targets per category and track week-over-week deltas.",
    ]
    return "\n".join(lines)

from typing import Optional, List, Dict
import io
from docx import Document

def _docx_from_report(
    company: str,
    days: int,
    category: Optional[str],
    activities: List[Dict],
    report_text: Optional[str],
) -> io.BytesIO:
    doc = Document()
    doc.add_heading(f"{company} — ESG/SDG Report", 0)

    sub = f"Period: last {days} days"
    if category:
        sub += f" | Category: {category}"
    doc.add_paragraph(sub)

    # Activity summary table
    doc.add_heading("Activity Summary", level=1)
    tbl = doc.add_table(rows=1, cols=4)
    hdr = tbl.rows[0].cells
    hdr[0].text = "Activity"
    hdr[1].text = "Description"
    hdr[2].text = "Persons"
    hdr[3].text = "Contributions"

    for a in activities or []:
        row = tbl.add_row().cells
        row[0].text = str(a.get("label", ""))
        row[1].text = str(a.get("action", ""))
        row[2].text = str(a.get("persons", 0))
        row[3].text = str(round(float(a.get("contributions", 0) or 0)))

    # Narrative
    doc.add_paragraph("") 
    doc.add_heading("Narrative Report", level=1)

    narrative = (report_text or "").strip()
    if not narrative:
        narrative = (
            "No AI narrative was generated for this period. "
            "This can happen if the language model API key is missing or the quota was exceeded."
        )

    for para in narrative.splitlines():
        doc.add_paragraph(para if para.strip() else "")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@router.get("/report/sdg.docx")
def download_sdg_docx(
    company: str = "Organization",
    days: int = 30,
    category: Optional[str] = None,
    model: Optional[str] = None,
):
    """
    Generate a Word (.docx) SDG/ESG report from activity_log using an OpenAI model.
    """
    activities = _fetch_activity_rollup(days=days, category=category)
    prompt = _build_prompt(company=company, activities=activities)
    report_text = _call_llm(prompt, model=model)
    buf = _docx_from_report(company, days, category, activities, report_text)

    ts = datetime.utcnow().strftime("%Y%m%d")
    filename = f"{company.replace(' ', '_')}_ESG_SDG_Report_{ts}.docx"
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )
# -------- NEW: category breakdown (for pies) --------
@router.get("/categories/summary")
def categories_summary(days: int = 30):
    cat_expr = _cat_case_expr("al")
    with engine.begin() as conn:
        where_sql, params = _where_time_and_category(days, None)
        sql = f"""
            SELECT {cat_expr} AS category,
                   COUNT(DISTINCT al.user_id)::int AS users,
                   COUNT(*)::int AS completions,
                   COALESCE(SUM(al.reward),0)::float AS points
              FROM public.activity_log al
             WHERE {where_sql}
             GROUP BY category
        """
        rows = conn.execute(text(sql), params).mappings().all()

    # ensuring always return all 6 categories (missing -> zeros)
    by_cat = {r["category"]: {"users": int(r["users"]), "completions": int(r["completions"]), "points": float(r["points"])} for r in rows}
    return [
        {"category": c, **by_cat.get(c, {"users": 0, "completions": 0, "points": 0.0})}
        for c in CAT_ORDER
    ]

# -------- NEW: category mini trends (2 buckets) --------
@router.get("/categories/trends")
def categories_trends(bucket: str = Query("day", pattern="^(day|week|month)$")):
    unit = {"day": "1 day", "week": "1 week", "month": "1 month"}[bucket]
    cat_expr = _cat_case_expr("al")

    # last two bucket windows (prev, curr)
    sql = f"""
    WITH buckets AS (
       SELECT date_trunc('{bucket}', now()) AS b0,
              date_trunc('{bucket}', now()) - INTERVAL '{unit}' AS b1
    ),
    agg AS (
       SELECT {cat_expr} AS category,
              date_trunc('{bucket}', al.ts) AS b,
              COUNT(*)::int AS c
         FROM public.activity_log al
        WHERE al.ts >= (SELECT b1 FROM buckets)
          AND al.ts <  (SELECT b0 FROM buckets) + INTERVAL '{unit}'
        GROUP BY 1,2
    )
    SELECT category,
           COALESCE(MAX(c) FILTER (WHERE b=(SELECT b1 FROM buckets)), 0) AS prev,
           COALESCE(MAX(c) FILTER (WHERE b=(SELECT b0 FROM buckets)), 0) AS curr
      FROM agg
     GROUP BY category
    """
    with engine.begin() as conn:
        rows = conn.execute(text(sql)).mappings().all()

    by_cat = {r["category"]: [int(r["prev"]), int(r["curr"])] for r in rows}
    return [
        {"id": c, "label": c, "trend": by_cat.get(c, [0, 0])}
        for c in CAT_ORDER
    ]
# Order we want categories to appear in responses / charts
CAT_ORDER: List[str] = ["donate", "volunteer", "advocate", "wellness", "recycle", "wildlife"]

def _cat_case_expr(alias: str = "al") -> str:
    """
    Returns a SQL CASE expression that maps action codes to our 6 category ids.
    Example output used in SELECTs:
        CASE
          WHEN al.action LIKE 'CAT01_%' THEN 'donate'
          WHEN al.action LIKE 'CAT02_%' THEN 'volunteer'
          ...
          ELSE 'unknown'
        END
    """
    return (
        "CASE "
        f"WHEN {alias}.action LIKE 'CAT01_%' THEN 'donate' "
        f"WHEN {alias}.action LIKE 'CAT02_%' THEN 'volunteer' "
        f"WHEN {alias}.action LIKE 'CAT03_%' THEN 'advocate' "
        f"WHEN {alias}.action LIKE 'CAT04_%' THEN 'wellness' "
        f"WHEN {alias}.action LIKE 'CAT05_%' THEN 'recycle' "
        f"WHEN {alias}.action LIKE 'CAT06_%' THEN 'wildlife' "
        "ELSE 'unknown' END"
    )
